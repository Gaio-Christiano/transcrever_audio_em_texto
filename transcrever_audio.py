# -----------------------------------------------------------------------------
# Passo 1: Importar todas as bibliotecas necessárias
# -----------------------------------------------------------------------------

import os  # Para interagir com o sistema operacional, como manipular nomes de arquivos.
import whisper  # A biblioteca de transcrição da OpenAI.
import torch  # Biblioteca principal de machine learning (PyTorch).
from pyannote.audio import Pipeline  # A biblioteca para identificar quem está falando (diarização).
from pydub import AudioSegment  # Para carregar e converter arquivos de áudio.
import tkinter as tk  # Para criar a janela de seleção de arquivo.
from tkinter import filedialog  # Especificamente para a caixa de diálogo de seleção de arquivo.
from docx import Document  # Para criar o arquivo Word (.docx).
from fpdf import FPDF  # Para criar o arquivo PDF.
import datetime  # Para formatar os tempos de início e fim.
import subprocess # Para executar comandos externos, como o FFmpeg.

# -----------------------------------------------------------------------------
# Passo 2: Configurações Iniciais e Constantes
# -----------------------------------------------------------------------------

# !!! IMPORTANTE !!!
# Cole aqui o seu token de acesso do Hugging Face.
# Este token é necessário para baixar e usar o modelo de identificação de locutores.
HF_TOKEN = "COLE_SEU_TOKEN_DO_HUGGING_FACE_AQUI"

# Verifica se um token foi inserido. Se não, o programa não pode continuar.
if HF_TOKEN == "COLE_SEU_TOKEN_DO_HUGGING_FACE_AQUI":
    print("ERRO: Por favor, insira seu token de acesso do Hugging Face na variável 'HF_TOKEN' no código.")
    exit() # Encerra o script se o token não for fornecido.

# Define se o processamento deve usar a GPU (mais rápido) ou a CPU.
# O torch.cuda.is_available() verifica se uma GPU compatível está instalada.
DEVICE = "cuda" if torch.cuda.is_available() else "cpu"
print(f"Usando dispositivo: {DEVICE}") # Informa ao usuário qual dispositivo será usado.

# -----------------------------------------------------------------------------
# Passo 3: Funções para Salvar o Resultado
# -----------------------------------------------------------------------------

def format_timestamp(seconds):
    """Converte segundos em um formato de tempo H:M:S,ms."""
    # Cria um objeto timedelta a partir dos segundos.
    td = datetime.timedelta(seconds=seconds)
    # Formata a string de tempo.
    return str(td)

def save_as_docx(transcript_data, output_filename):
    """Salva a transcrição formatada em um arquivo Word (.docx)."""
    # Cria um novo documento Word em branco.
    doc = Document()
    # Adiciona um título ao documento.
    doc.add_heading(f'Transcrição de {os.path.basename(output_filename)}', level=0)
    
    # Itera sobre cada entrada da transcrição (cada fala de um locutor).
    for entry in transcript_data:
        # Formata o tempo de início e fim.
        start_time = format_timestamp(entry['start'])
        end_time = format_timestamp(entry['end'])
        
        # Cria a linha do roteiro, como "NARRADOR 1 (0:00:10.500 --> 0:00:15.200): Texto da fala..."
        line = f"{entry['speaker']} ({start_time} --> {end_time}):\n{entry['text']}\n"
        # Adiciona a linha ao documento.
        doc.add_paragraph(line)
        
    # Salva o documento com o nome de arquivo fornecido.
    doc.save(f"{output_filename}.docx")
    print(f"Arquivo Word salvo como: {output_filename}.docx")

def save_as_pdf(transcript_data, output_filename):
    """Salva a transcrição formatada em um arquivo PDF."""
    # Cria um novo objeto PDF.
    pdf = FPDF()
    # Adiciona uma página.
    pdf.add_page()
    # Define a fonte. 'DejaVu' é usada para suportar caracteres especiais como acentos.
    pdf.add_font('DejaVu', '', 'DejaVuSans.ttf', uni=True)
    pdf.set_font('DejaVu', '', 10)
    
    # Adiciona o título ao PDF.
    pdf.cell(0, 10, f'Transcrição de {os.path.basename(output_filename)}', ln=True, align='C')
    
    # Itera sobre cada entrada da transcrição.
    for entry in transcript_data:
        start_time = format_timestamp(entry['start'])
        end_time = format_timestamp(entry['end'])
        
        # Cria a linha do roteiro.
        line = f"{entry['speaker']} ({start_time} --> {end_time}):\n{entry['text']}\n"
        # Adiciona a linha ao PDF usando multi_cell para quebras de linha automáticas.
        pdf.multi_cell(0, 5, line.encode('latin-1', 'replace').decode('latin-1'))
        
    # Salva o arquivo PDF.
    pdf.output(f"{output_filename}.pdf")
    print(f"Arquivo PDF salvo como: {output_filename}.pdf")


# -----------------------------------------------------------------------------
# Passo 4: Função Principal de Processamento
# -----------------------------------------------------------------------------

def process_audio(input_mp3_path):
    """Função principal que executa todo o processo de transcrição e diarização."""
    
    print("Iniciando o processamento do áudio. Isso pode levar muito tempo para arquivos grandes...")

    # Define os nomes dos arquivos de saída com base no nome do arquivo de entrada.
    base_filename = os.path.splitext(os.path.basename(input_mp3_path))[0]
    output_wav_path = f"{base_filename}_converted.wav"

    """# --- ETAPA A: Conversão de MP3 para WAV ---
    print(f"1/5 - Convertendo '{input_mp3_path}' para o formato WAV...")
    try:
        # Carrega o arquivo MP3 usando pydub.
        audio = AudioSegment.from_mp3(input_mp3_path)
        # Converte para mono (1 canal) e define a taxa de amostragem para 16kHz, ideal para os modelos.
        audio = audio.set_channels(1).set_frame_rate(16000)
        # Exporta (salva) o áudio convertido como um arquivo WAV.
        audio.export(output_wav_path, format="wav")
    except Exception as e:
        print(f"Erro ao converter o áudio. Verifique se o FFmpeg está instalado corretamente. Erro: {e}")
        return"""

    # --- ETAPA A: Conversão de MP3 para WAV (Método Robusto) ---
    print(f"1/5 - Convertendo '{input_mp3_path}' para o formato WAV usando FFmpeg diretamente...")
    try:
        # Constrói o comando do FFmpeg para converter o áudio.
        # -i: arquivo de entrada
        # -ar 16000: define a taxa de amostragem para 16kHz
        # -ac 1: define para 1 canal de áudio (mono)
        # -y: sobrescreve o arquivo de saída se ele já existir
        command = [
            'ffmpeg',
            '-i', input_mp3_path,
            '-ar', '16000',
            '-ac', '1',
            '-y',
            output_wav_path
        ]
        
        # Executa o comando FFmpeg. O 'check=True' fará o script parar se o FFmpeg falhar.
        subprocess.run(command, check=True, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
        
    except subprocess.CalledProcessError:
        # Captura um erro se o FFmpeg falhar na conversão.
        print(f"Erro ao converter o áudio com FFmpeg. Verifique se o FFmpeg está instalado e acessível no PATH do sistema.")
        return
    except Exception as e:
        # Captura outros erros inesperados.
        print(f"Um erro inesperado ocorreu durante a conversão: {e}")
        return

    # --- ETAPA B: Identificação do Locutor (Diarização) ---
    print("2/5 - Identificando os locutores no áudio (diarização)...")
    # Carrega o pipeline de diarização pré-treinado da pyannote.audio.
    # Usa o token de acesso para autenticação.
    diarization_pipeline = Pipeline.from_pretrained(
        "pyannote/speaker-diarization-3.1",
        use_auth_token=HF_TOKEN
    ).to(torch.device(DEVICE))
    
    # Executa o pipeline no arquivo WAV. O resultado contém os segmentos de fala de cada locutor.
    diarization_result = diarization_pipeline(output_wav_path)

    # --- ETAPA C: Transcrição do Áudio (Fala para Texto) ---
    print("3/5 - Transcrevendo o áudio para texto com o Whisper...")
    # Carrega o modelo Whisper. "base" é um bom equilíbrio entre velocidade e precisão.
    # Para maior precisão (e maior tempo de processamento), use "medium" ou "large".
    whisper_model = whisper.load_model("base", device=DEVICE)
    # Executa a transcrição no arquivo WAV, pedindo os timestamps de cada palavra.
    transcription_result = whisper_model.transcribe(output_wav_path, word_timestamps=True)

    # --- ETAPA D: Combinando Resultados da Diarização e Transcrição ---
    print("4/5 - Combinando a transcrição com a identificação dos locutores...")
    
    # Mapeia as palavras transcritas para seus respectivos locutores.
    word_speakers = []
    # Itera sobre cada segmento de fala identificado pela diarização.
    for segment, _, speaker in diarization_result.itertracks(yield_label=True):
        # Itera sobre cada palavra transcrita pelo Whisper.
        for word in transcription_result['segments'][0]['words']:
            # Verifica se o tempo da palavra está dentro do segmento do locutor.
            if segment.start <= word['start'] < segment.end:
                # Se estiver, associa a palavra ao locutor.
                word_speakers.append({'word': word['word'], 'start': word['start'], 'end': word['end'], 'speaker': speaker})

    # Agrupa as palavras em frases contínuas para cada locutor.
    final_transcript = []
    if word_speakers:
        # Inicia com a primeira palavra.
        current_entry = {
            'speaker': f"NARRADOR {int(word_speakers[0]['speaker'].split('_')[1]) + 1}", # Renomeia 'SPEAKER_00' para 'NARRADOR 1'
            'text': word_speakers[0]['word'],
            'start': word_speakers[0]['start'],
            'end': word_speakers[0]['end']
        }

        # Itera sobre as palavras restantes.
        for i in range(1, len(word_speakers)):
            current_word = word_speakers[i]
            prev_word = word_speakers[i-1]
            
            speaker_id = f"NARRADOR {int(current_word['speaker'].split('_')[1]) + 1}"
            
            # Se o locutor for o mesmo e o tempo entre as palavras for curto, agrupa.
            if current_word['speaker'] == prev_word['speaker'] and (current_word['start'] - prev_word['end']) < 0.5:
                current_entry['text'] += ' ' + current_word['word']
                current_entry['end'] = current_word['end']
            else:
                # Se o locutor mudar ou houver uma pausa longa, salva a entrada anterior e começa uma nova.
                final_transcript.append(current_entry)
                current_entry = {
                    'speaker': speaker_id,
                    'text': current_word['word'],
                    'start': current_word['start'],
                    'end': current_word['end']
                }
        # Adiciona a última entrada à lista.
        final_transcript.append(current_entry)
    
    # --- ETAPA E: Salvando os Arquivos de Saída ---
    print(f"5/5 - Salvando os arquivos de saída com o nome base '{base_filename}'...")
    save_as_docx(final_transcript, base_filename)
    save_as_pdf(final_transcript, base_filename)

    # Limpeza: remove o arquivo WAV temporário.
    os.remove(output_wav_path)
    print("\nProcesso concluído com sucesso!")


# -----------------------------------------------------------------------------
# Passo 5: Ponto de Entrada do Script
# -----------------------------------------------------------------------------

if __name__ == "__main__":
    # Cria uma janela raiz do tkinter, mas a oculta.
    root = tk.Tk()
    root.withdraw()

    print("Por favor, selecione o arquivo de áudio MP3 que você deseja transcrever.")
    
    # Abre a caixa de diálogo para o usuário selecionar um arquivo.
    # O usuário pode selecionar arquivos .mp3.
    file_path = filedialog.askopenfilename(
        title="Selecione o arquivo de áudio MP3",
        filetypes=(("Arquivos MP3", "*.mp3"), ("Todos os arquivos", "*.*"))
    )

    # Se um arquivo foi selecionado (o caminho não está vazio).
    if file_path:
        try:
            # Chama a função principal de processamento.
            process_audio(file_path)
        except Exception as e:
            # Captura qualquer erro inesperado durante o processo.
            print(f"\nOcorreu um erro inesperado: {e}")
    else:
        # Se nenhum arquivo foi selecionado.

        print("Nenhum arquivo selecionado. O programa será encerrado.")
